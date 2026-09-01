import os
import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from datetime import datetime
import ai_engine as ai
import utils as util

# ==============================================================================
# 1. FUNÇÕES AUXILIARES TÉCNICAS E SANITIZAÇÃO XML SOBERANA
# ==============================================================================

def sanitizar_xml_str(texto):
    """
    SOSA V2026 - VACINA ANTI-ERRO LXML/WORD:
    Remove caracteres de controle nulos e invisíveis do XML (padrão W3C XML 1.0).
    Preserva tabulações (\t), quebras de linha (\n), retornos (\r) e todos os
    caracteres imprimíveis Unicode/ASCII (acentos, emojis e LaTeX $$).
    """
    if not texto or not isinstance(texto, str):
        return ""
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', texto)

def set_row_height(row, height_pt):
    """Define a altura mínima da linha da tabela para o cabeçalho não achatar"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pt * 20))) 
    trHeight.set(qn('w:hRule'), "atLeast")
    trPr.append(trHeight)

def set_cell_background(cell, fill_hex):
    """Aplica cor de fundo executiva em uma célula da tabela"""
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)
    except: pass

def helper_sosa_float(v):
    """Converte qualquer valor para float de forma imune a erros e textos anexados"""
    if v is None or str(v).strip() == "" or str(v).lower() == "nan": return 0.0
    try:
        s = str(v).strip().replace(" ", "").replace(",", ".")
        m = re.search(r'[-+]?\d*\.?\d+', s)
        if m:
            return float(m.group(0))
        return 0.0
    except:
        return 0.0

def converter_latex_para_texto_word(texto):
    """
    SOSA V2026 - VACINA INVIOLÁVEL DO CIFRÃO E DO CONTROL CHAR:
    1. Higieniza o texto contra caracteres de controle XML inválidos.
    2. Preserva 100% dos marcadores $$ para compatibilidade com o Google Docs Apps Script.
    3. Corrige R\$ para R$ (dinheiro) e \% para %.
    """
    if not texto or not isinstance(texto, str): return ""
    t = sanitizar_xml_str(texto)
    t = t.replace(r'R\$', 'R$').replace(r'R \$', 'R$')
    t = t.replace(r'\$', '$').replace(r'\%', '%')
    return t.strip()

def adicionar_texto_formatado(paragraph, texto, cor_rgb=None, tamanho_pt=None):
    """Converte padrões **texto** em negrito real preservando expressões matemáticas $$"""
    if not texto: return
    texto_limpo = converter_latex_para_texto_word(texto)
    texto_limpo = texto_limpo.replace("➔", "").replace("->", "→").strip()
    
    partes = re.split(r'(\*\*.*?\*\*)', texto_limpo)
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**'):
            run = paragraph.add_run(parte.replace('**', ''))
            run.bold = True
        else:
            run = paragraph.add_run(parte)
            
        if cor_rgb:
            run.font.color.rgb = cor_rgb
        if tamanho_pt:
            run.font.size = Pt(tamanho_pt)

def adicionar_box_imagem_word(doc, legenda_prompt="ESPAÇO PARA ILUSTRAÇÃO / DESENHO"):
    """Cria uma moldura visual elegante para o prompt de ilustração no Word sem poluir o layout"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    set_cell_background(cell, "F8FAFC")
    set_row_height(table.rows[0], 45)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    run_icon = p.add_run("🖼️ [ILUSTRAÇÃO TÉCNICA A4]: ")
    run_icon.font.bold = True
    run_icon.font.size = Pt(9.0)
    run_icon.font.color.rgb = RGBColor(41, 98, 255)
    
    run_desc = p.add_run(f"{sanitizar_xml_str(str(legenda_prompt)).strip()}")
    run_desc.font.size = Pt(8.5)
    run_desc.font.italic = True
    run_desc.font.color.rgb = RGBColor(71, 85, 105)
    
    doc.add_paragraph()

def renderizar_tabela_markdown_no_word(doc, linhas_tabela):
    """
    SOSA V2026: Converte linhas brutas de tabela Markdown (| Coluna 1 | Coluna 2 |)
    em uma tabela nativa oficial do Word com suporte a negrito e LaTeX formatado.
    """
    if not linhas_tabela: return
    
    linhas_processadas = []
    for linha in linhas_tabela:
        linha_limpa = linha.strip()
        if not linha_limpa: continue
        if re.match(r'^\|?\s*:?-+:?\s*(\|?\s*:?-+:?\s*)*\|?$', linha_limpa):
            continue
        celulas = [c.strip() for c in linha_limpa.strip('|').split('|')]
        if any(c for c in celulas):
            linhas_processadas.append(celulas)
            
    if not linhas_processadas: return
    
    num_cols = max(len(r) for r in linhas_processadas)
    num_rows = len(linhas_processadas)
    
    table_word = doc.add_table(rows=num_rows, cols=num_cols)
    table_word.style = 'Table Grid'
    table_word.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for r_idx, row_data in enumerate(linhas_processadas):
        set_row_height(table_word.rows[r_idx], 18)
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = table_word.cell(r_idx, c_idx)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                
                if r_idx == 0:
                    set_cell_background(cell, "003366")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    adicionar_texto_formatado(p, cell_text, cor_rgb=RGBColor(255, 255, 255), tamanho_pt=8.5)
                    for r in p.runs: r.bold = True
                else:
                    if r_idx % 2 == 0:
                        set_cell_background(cell, "F8FAFC")
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    adicionar_texto_formatado(p, cell_text, tamanho_pt=8.5)
    doc.add_paragraph()

def renderizar_conteudo_com_tabelas(doc, texto_bruto):
    """Processa o texto alternando entre parágrafos normais e tabelas Markdown nativas."""
    linhas = sanitizar_xml_str(str(texto_bruto)).split('\n')
    buffer_tabela = []
    em_tabela = False
    
    for linha in linhas:
        l_s = linha.strip()
        
        if l_s.startswith('|') and l_s.endswith('|'):
            em_tabela = True
            buffer_tabela.append(l_s)
            continue
        else:
            if em_tabela:
                renderizar_tabela_markdown_no_word(doc, buffer_tabela)
                buffer_tabela = []
                em_tabela = False
                
        if not l_s: continue
        
        if "[" in l_s and "PROMPT IMAGEM" in l_s.upper():
            desc_p = re.sub(r'\[\s*PROMPT IMAGEM:\s*|\s*\]', '', l_s, flags=re.IGNORECASE)
            adicionar_box_imagem_word(doc, desc_p)
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15

        if l_s.upper().startswith("CONCEITOS CHAVE") or l_s.upper().startswith("DICA"):
            run_c = p.add_run("📌 " + l_s)
            run_c.bold = True
            run_c.font.size = Pt(9.5)
            run_c.font.color.rgb = RGBColor(0, 51, 102)
            continue

        if l_s.upper().startswith("QUESTÃO") or l_s.upper().startswith("ITEM"):
            match = re.match(r"^(QUEST[AÃ]O\s*\d+|ITEM\s*\d+)(\s*\(.*?\))?([\s\.\-\:]+)(.*)", l_s, re.IGNORECASE)
            if match:
                rotulo = f"{match.group(1).upper()}{match.group(2) if match.group(2) else ''}{match.group(3)}"
                run_r = p.add_run(rotulo)
                run_r.bold = True
                run_r.font.size = Pt(10.0)
                run_r.font.color.rgb = RGBColor(0, 51, 102)
                adicionar_texto_formatado(p, match.group(4).strip())
                continue

        if re.match(r'^[A-E][\)\.]', l_s):
            p.paragraph_format.left_indent = Inches(0.15)
            letra_match = re.match(r'^([A-E][\)\.])(.*)', l_s)
            if letra_match:
                run_letra = p.add_run(letra_match.group(1))
                run_letra.bold = True
                adicionar_texto_formatado(p, letra_match.group(2))
                continue
        
        adicionar_texto_formatado(p, l_s)
        
    if em_tabela and buffer_tabela:
        renderizar_tabela_markdown_no_word(doc, buffer_tabela)

def configurar_cabecalho_mestre(doc, info, tipo_label, mostrar_nota=False):
    """Gera o cabeçalho executivo oficial da Prefeitura e Escola de Itabuna perfeitamente calibrado para A4"""
    table = doc.add_table(rows=3, cols=5)
    table.style = 'Table Grid'
    
    widths = [Inches(0.75), Inches(3.0), Inches(0.9), Inches(1.0), Inches(1.8)]
    for i, w in enumerate(widths): 
        table.columns[i].width = w

    for row in table.rows:
        set_row_height(row, 22)

    # LINHA 0: LOGO, ESCOLA E TRIMESTRE
    c_logo = table.cell(0, 0).merge(table.cell(2, 0))
    c_escola = table.cell(0, 1).merge(table.cell(0, 3))
    p_esc = c_escola.paragraphs[0]
    p_esc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_esc = p_esc.add_run("ESCOLA MUNICIPAL FLAVIO JOSE SIMOES COSTA")
    run_esc.bold = True
    run_esc.font.size = Pt(10)
    
    c_trim = table.cell(0, 4)
    c_trim.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_background(c_trim, "F1F5F9")
    p_tr = c_trim.paragraphs[0]
    p_tr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tr = p_tr.add_run(sanitizar_xml_str(str(info.get('trimestre', 'I Trimestre'))))
    run_tr.font.bold = True
    run_tr.font.size = Pt(9.5)

    # LINHA 1: ALUNO
    if mostrar_nota:
        c_aluno = table.cell(1, 1).merge(table.cell(1, 3))
        c_aluno.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c_aluno.paragraphs[0].add_run("ESTUDANTE: __________________________________________________").font.size = Pt(9.5)
        
        c_nota = table.cell(1, 4)
        c_nota.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(c_nota, "F8FAFC")
        p_n = c_nota.paragraphs[0]
        p_n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_n.add_run("NOTA: ________").font.bold = True
    else:
        c_aluno = table.cell(1, 1).merge(table.cell(1, 4))
        c_aluno.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c_aluno.paragraphs[0].add_run("ESTUDANTE: ________________________________________________________________").font.size = Pt(9.5)

    # LINHA 2: PROFESSOR, TURMA, DATA E TIPO
    table.cell(2, 1).paragraphs[0].add_run("PROF: Ronaldo Gomes").font.size = Pt(9)
    table.cell(2, 2).paragraphs[0].add_run(f"TURMA: {sanitizar_xml_str(str(info.get('ano', '6º')))}").font.size = Pt(9)
    table.cell(2, 3).paragraphs[0].add_run("DATA:    /    /").font.size = Pt(9)
    
    c_tipo = table.cell(2, 4)
    c_tipo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_background(c_tipo, "2962FF")
    p_tipo = c_tipo.paragraphs[0]
    p_tipo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tipo = p_tipo.add_run(sanitizar_xml_str(str(tipo_label)))
    run_tipo.font.bold = True
    run_tipo.font.size = Pt(9)
    run_tipo.font.color.rgb = RGBColor(255, 255, 255)

    logo_path = "logo_escola.png" if os.path.exists("logo_escola.png") else "logo.png"
    if os.path.exists(logo_path):
        c_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c_logo.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try: p.add_run().add_picture(logo_path, width=Inches(0.65))
        except: pass
    return table

# ==============================================================================
# 2. MATERIAL DO ALUNO REGULAR
# ==============================================================================
def gerar_docx_aluno_v24(titulo_doc, conteudo, info):
    file_stream = io.BytesIO()
    doc = Document()
    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Inches(0.4), Inches(0.4)
    section.left_margin, section.right_margin = Inches(0.4), Inches(0.4)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    configurar_cabecalho_mestre(doc, info, "ATIVIDADE DE SALA", mostrar_nota=False)
    doc.add_paragraph()

    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    sectPr = new_section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '420')

    renderizar_conteudo_com_tabelas(doc, conteudo)

    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ==============================================================================
# 3. MATERIAL PEI ADAPTADO (NÍVEIS 1 E 2)
# ==============================================================================
def gerar_docx_pei_v25(titulo_doc, conteudo, info):
    file_stream = io.BytesIO()
    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin, section.bottom_margin = Inches(0.4), Inches(0.4)
        section.left_margin, section.right_margin = Inches(0.4), Inches(0.4)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        conteudo_limpo = sanitizar_xml_str(str(conteudo)).strip()
        conteudo_limpo = re.sub(r'^(?:Olá|Como especialista|Como profissional|Prezado|Segue).*?\n\n', '', conteudo_limpo, flags=re.IGNORECASE | re.DOTALL).strip()

        label_pei = "AVALIAÇÃO ADAPTADA (PEI NÍVEL 1)" if "N1" in titulo_doc.upper() or "NIVEL_1" in titulo_doc.upper() else "AVALIAÇÃO ADAPTADA (PEI NÍVEL 2)"
        configurar_cabecalho_mestre(doc, info, label_pei, mostrar_nota=True)
        doc.add_paragraph()

        num_total_q = len(re.findall(r'(?i)(?:QUEST[AÃ]O\s*(?:PEI\s*)?|Q)\s*\d+', conteudo_limpo))
        if num_total_q == 0: 
            num_total_q = int(helper_sosa_float(info.get('qtd', 10)))

        adicionar_cartao_resposta_fiducial_word(doc, num_total_q, is_pei=True)
        doc.add_paragraph()

        new_section = doc.add_section(WD_SECTION.CONTINUOUS)
        sectPr = new_section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '420')

        renderizar_conteudo_com_tabelas(doc, conteudo_limpo)

        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        file_stream = io.BytesIO()
        err_doc = Document()
        err_doc.add_paragraph(f"ERRO NO EXPORTER PEI: {sanitizar_xml_str(str(e))}")
        err_doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

# ==============================================================================
# 4. GUIA DO PROFESSOR
# ==============================================================================
def gerar_docx_professor_v25(titulo_doc, conteudo, info):
    file_stream = io.BytesIO()
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Inches(0.4), Inches(0.4)
    section.left_margin, section.right_margin = Inches(0.4), Inches(0.4)

    header_table = doc.add_table(rows=2, cols=3)
    header_table.style = 'Table Grid'
    c_tit = header_table.cell(0, 0).merge(header_table.cell(0, 2))
    set_cell_background(c_tit, "003366")
    run_tit = c_tit.paragraphs[0].add_run("GUIA DE MEDIAÇÃO, DESCRITORES SAEB E DISTRATORES")
    run_tit.font.bold, run_tit.font.size = True, Pt(11)
    run_tit.font.color.rgb = RGBColor(255, 255, 255)
    c_tit.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    header_table.cell(1, 0).paragraphs[0].add_run(f"ANO: {sanitizar_xml_str(str(info.get('ano', '')))}").font.size = Pt(9)
    header_table.cell(1, 1).paragraphs[0].add_run(f"SEMANA: {sanitizar_xml_str(str(info.get('semana', '')))}").font.size = Pt(9)
    header_table.cell(1, 2).paragraphs[0].add_run(f"TRIMESTRE: {sanitizar_xml_str(str(info.get('trimestre', 'I')))}").font.size = Pt(9)
    for row in header_table.rows: set_row_height(row, 20)
    doc.add_paragraph()

    renderizar_conteudo_com_tabelas(doc, conteudo)

    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ==============================================================================
# 5. PROVA OFICIAL & CARTÃO OMR FIDUCIAL (PADRÃO ENEM / SAEB)
# ==============================================================================

def adicionar_cartao_resposta_fiducial_word(doc, num_total_q, is_pei=False):
    """Cartão-Resposta com Colunas Travadas em 1,0 cm (Cm(1.0)) e 4 Marcadores Fiduciais (■)."""
    container_table = doc.add_table(rows=3, cols=3)
    container_table.style = 'Table Grid'
    container_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    container_table.columns[0].width = Inches(0.45)
    container_table.columns[1].width = Inches(6.5)
    container_table.columns[2].width = Inches(0.45)

    set_row_height(container_table.rows[0], 28)
    set_row_height(container_table.rows[2], 28)

    for r_idx, c_idx in [(0, 0), (0, 2), (2, 0), (2, 2)]:
        c = container_table.cell(r_idx, c_idx)
        set_cell_background(c, "000000")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("■")
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(16)

    c_title = container_table.cell(0, 1)
    set_cell_background(c_title, "F1F5F9")
    p_t = c_title.paragraphs[0]
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("CARTÃO-RESPOSTA OFICIAL (FOLHA DE RESPOSTAS)")
    r_t.font.bold = True
    r_t.font.size = Pt(10.5)

    c_foot = container_table.cell(2, 1)
    set_cell_background(c_foot, "F8FAFC")
    p_f = c_foot.paragraphs[0]
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_f.add_run("▲ MANTENHA O PAPEL RETO • PREENCHA TOTALMENTE OS CÍRCULOS COM CANETA PRETA OU AZUL ▲")
    r_f.font.size = Pt(8.0)
    r_f.font.bold = True
    r_f.font.color.rgb = RGBColor(100, 116, 139)

    c_grid = container_table.cell(1, 1)
    col_count = 4 if is_pei else 6
    headers = ["Q", "A", "B", "C"] if is_pei else ["Q", "A", "B", "C", "D", "E"]
    
    if num_total_q <= 10:
        gab_grid = c_grid.add_table(rows=num_total_q + 1, cols=col_count)
        gab_grid.style = 'Table Grid'
        gab_grid.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, col in enumerate(gab_grid.columns):
            if i == 0: col.width = Cm(0.8)
            else: col.width = Cm(1.0)
        
        for i, lab in enumerate(headers):
            c = gab_grid.cell(0, i)
            set_cell_background(c, "2962FF" if i==0 else "E2E8F0")
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_h = p.add_run(lab)
            r_h.font.bold = True
            r_h.font.size = Pt(9.5)
            if i == 0: r_h.font.color.rgb = RGBColor(255, 255, 255)
            
        for r in range(1, num_total_q + 1):
            set_row_height(gab_grid.rows[r], 22)
            c_q = gab_grid.cell(r, 0)
            set_cell_background(c_q, "F1F5F9")
            p_q = c_q.paragraphs[0]
            p_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_q_num = p_q.add_run(f"{r:02d}")
            r_q_num.font.size = Pt(9.5)
            r_q_num.font.bold = True
            
            for col in range(1, col_count):
                c_b = gab_grid.cell(r, col)
                p_b = c_b.paragraphs[0]
                p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_b.add_run("○").font.size = Pt(13)
    else:
        half = (num_total_q + 1) // 2
        double_cols = col_count * 2
        gab_grid = c_grid.add_table(rows=half + 1, cols=double_cols)
        gab_grid.style = 'Table Grid'
        gab_grid.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, col in enumerate(gab_grid.columns):
            if i % col_count == 0: col.width = Cm(0.8)
            else: col.width = Cm(1.0)
        
        headers_double = headers + headers
        for i, lab in enumerate(headers_double):
            c = gab_grid.cell(0, i)
            is_q_col = (i % col_count == 0)
            set_cell_background(c, "2962FF" if is_q_col else "E2E8F0")
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_h = p.add_run(lab)
            r_h.font.bold = True
            r_h.font.size = Pt(9.0)
            if is_q_col: r_h.font.color.rgb = RGBColor(255, 255, 255)

        for r in range(1, num_total_q + 1):
            row_idx = r if r <= half else r - half
            col_offset = 0 if r <= half else col_count
            
            if row_idx < len(gab_grid.rows):
                set_row_height(gab_grid.rows[row_idx], 20)
                c_q = gab_grid.cell(row_idx, col_offset)
                set_cell_background(c_q, "F1F5F9")
                p_q = c_q.paragraphs[0]
                p_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_q = p_q.add_run(f"{r:02d}")
                r_q.font.size = Pt(9.0)
                r_q.font.bold = True
                
                for col in range(1, col_count):
                    c_b = gab_grid.cell(row_idx, col_offset + col)
                    p_b = c_b.paragraphs[0]
                    p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_b.add_run("○").font.size = Pt(12)

def adicionar_box_calculo_discursivo_word(doc):
    """Cria uma caixa pautada oficial para resolução de cálculo e resposta final do aluno."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(7.2)
    
    set_cell_background(cell, "FAFAFA")
    set_row_height(table.rows[0], 90)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    run_icon = p.add_run("✍️ MEMÓRIA DE CÁLCULO E RESPOSTA FINAL DO ESTUDANTE:\n")
    run_icon.font.bold = True
    run_icon.font.size = Pt(8.5)
    run_icon.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph()

def renderizar_conteudo_discursivo_com_caixas(doc, texto_bruto):
    """Renderiza questões abertas com caixas pautadas de resolução para cada item."""
    linhas = sanitizar_xml_str(str(texto_bruto)).split('\n')
    buffer_tabela = []
    em_tabela = False
    
    for linha in linhas:
        l_s = linha.strip()
        
        if l_s.startswith('|') and l_s.endswith('|'):
            em_tabela = True
            buffer_tabela.append(l_s)
            continue
        else:
            if em_tabela:
                renderizar_tabela_markdown_no_word(doc, buffer_tabela)
                buffer_tabela = []
                em_tabela = False
                
        if not l_s: continue
        
        if "[" in l_s and "PROMPT IMAGEM" in l_s.upper():
            desc_p = re.sub(r'\[\s*PROMPT IMAGEM:\s*|\s*\]', '', l_s, flags=re.IGNORECASE)
            adicionar_box_imagem_word(doc, desc_p)
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15

        if l_s.upper().startswith("QUESTÃO") or l_s.upper().startswith("ITEM"):
            match = re.match(r"^(QUEST[AÃ]O\s*\d+|ITEM\s*\d+)(\s*\(.*?\))?([\s\.\-\:]+)(.*)", l_s, re.IGNORECASE)
            if match:
                rotulo = f"{match.group(1).upper()}{match.group(2) if match.group(2) else ''}{match.group(3)}"
                run_r = p.add_run(rotulo)
                run_r.bold = True
                run_r.font.size = Pt(10.5)
                run_r.font.color.rgb = RGBColor(0, 51, 102)
                adicionar_texto_formatado(p, match.group(4).strip())
                adicionar_box_calculo_discursivo_word(doc)
                continue

        adicionar_texto_formatado(p, l_s)
        
    if em_tabela and buffer_tabela:
        renderizar_tabela_markdown_no_word(doc, buffer_tabela)

# ==============================================================================
# 5. GERADOR HÍBRIDO DE AVALIAÇÕES (REGULAR DISCURSIVA OU OMR / PEI LIMPO EM 2 COLUNAS)
# ==============================================================================
def gerar_docx_prova_v25(titulo_doc, conteudo_ia, info):
    """
    SOSA V2026.PRO_INFINITY - GERADOR HÍBRIDO DE AVALIAÇÕES
    - Provas Regulares Discursivas / 2ª Chamada: 1 coluna larga com caixas pautadas de cálculo.
    - Provas PEI Adaptadas: 2 colunas elegantes com 3 alternativas (A, B, C), SEM caixas deformadas.
    - Provas Regulares Objetivas: 2 colunas com Cartão OMR Fiducial (■).
    """
    file_stream = io.BytesIO()
    try:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        section = doc.sections[0]
        section.top_margin = section.bottom_margin = Inches(0.35)
        section.left_margin = section.right_margin = Inches(0.4)
        
        conteudo_ia_limpo = sanitizar_xml_str(str(conteudo_ia))
        is_pei_doc = any(x in str(titulo_doc).upper() for x in ["PEI", "ADAPTADA", "N1", "N2", "NIVEL_1", "NIVEL_2"])
        
        tipo_prova_info = str(info.get('tipo_prova', '')).upper()
        titulo_upper = str(titulo_doc).upper()
        
        is_discursiva = any(x in titulo_upper or x in tipo_prova_info for x in [
            "RECUPERAÇÃO", "RECUPERACAO", "2ª CHAMADA", "2A CHAMADA", "2ª_CHAMADA", "DISCURSIVA", "ABERTA"
        ]) and "FINAL" not in titulo_upper and not is_pei_doc

        corpo_bruto = ""
        if is_pei_doc:
            corpo_bruto = (ai.extrair_tag(conteudo_ia_limpo, "PEI_NIVEL_1") or 
                           ai.extrair_tag(conteudo_ia_limpo, "PEI") or 
                           ai.extrair_tag(conteudo_ia_limpo, "PEI_NIVEL_2") or 
                           ai.extrair_tag(conteudo_ia_limpo, "NIVEL_1"))
        else:
            corpo_bruto = ai.extrair_tag(conteudo_ia_limpo, "ALUNO") or ai.extrair_tag(conteudo_ia_limpo, "QUESTOES")

        if not corpo_bruto or len(corpo_bruto.strip()) < 10:
            match_primeira_q = re.search(r"(?i)QUESTÃO\s*\d+", conteudo_ia_limpo)
            if match_primeira_q:
                corpo_bruto = conteudo_ia_limpo[match_primeira_q.start():].strip()
            else:
                corpo_bruto = conteudo_ia_limpo.strip()

        corpo_bruto = sanitizar_xml_str(corpo_bruto)

        num_total_q = len(re.findall(r'(?i)QUESTÃO\s+\d+', corpo_bruto))
        if num_total_q == 0: 
            num_total_q = int(helper_sosa_float(info.get('qtd_questoes', info.get('qtd', 10))))
        
        if is_pei_doc:
            label_prova = "AVALIAÇÃO ADAPTADA (RECUPERAÇÃO PEI)" if "RECUPERA" in titulo_upper else "AVALIAÇÃO ADAPTADA (PEI)"
        elif is_discursiva:
            if "RECUPERAÇÃO" in titulo_upper or "RECUPERACAO" in titulo_upper:
                label_prova = "RECUPERAÇÃO PARALELA (AVALIAÇÃO DISCURSIVA)"
            else:
                label_prova = "AVALIAÇÃO DE SEGUNDA CHAMADA (DISCURSIVA)"
        else:
            label_prova = "AVALIAÇÃO DE MATEMÁTICA (ENEM/SAEB)"

        val_total_num = 10.0 if any(x in titulo_upper or x in tipo_prova_info for x in ["RECUPERAÇÃO", "RECUPERACAO", "REC_"]) else helper_sosa_float(info.get('valor', 4.0))
        if val_total_num == 0: val_total_num = 10.0

        val_q_calc = val_total_num / num_total_q if num_total_q > 0 else 1.0
        val_total_str = f"{val_total_num:.1f}"
        val_q_str = f"{val_q_calc:.2f}".replace(".", ",")
        
        info_cabecalho = info.copy()
        info_cabecalho['valor'] = val_total_str
        info_cabecalho['valor_questao'] = val_q_str

        # 1. CABEÇALHO MESTRE
        configurar_cabecalho_mestre(doc, info_cabecalho, label_prova, mostrar_nota=True)
        doc.add_paragraph()

        # 2. ORIENTAÇÕES DE EXECUÇÃO
        top_table = doc.add_table(rows=1, cols=1)
        top_table.style = 'Table Grid'
        top_table.columns[0].width = Inches(7.2)
        c_orient = top_table.cell(0, 0)
        set_cell_background(c_orient, "F8FAFC")
        
        p_tit = c_orient.paragraphs[0]
        r_tit_inst = p_tit.add_run("📋 ORIENTAÇÕES DE EXECUÇÃO:")
        r_tit_inst.bold = True
        r_tit_inst.font.size = Pt(9.5)
        r_tit_inst.font.color.rgb = RGBColor(0, 51, 102)
        
        if is_discursiva:
            orient_list = [
                f"Valor Total: {val_total_str} pontos (Escala 0 a 10) | Questões Discursivas.",
                "Apresente de forma clara a memória de cálculo em cada questão no espaço reservado.",
                "Respostas sem a devida demonstração matemática do cálculo não receberão pontuação integral.",
                "Utilize caneta esferográfica preta ou azul para a declaração da resposta final."
            ]
        elif is_pei_doc:
            orient_list = [
                f"Valor Total: {val_total_str} pontos (Escala 0 a 10) | Avaliação Adaptada.",
                "Leia atentamente as questões e marque a alternativa correta (A, B ou C).",
                "Consulte as dicas [PARA LEMBRAR] em cada questão para auxiliar na resolução."
            ]
        else:
            orient_list = [
                f"Valor Total: {val_total_str} pts | Valor por Questão: {val_q_str} pts.",
                "Preencha o Cartão-Resposta com caneta esferográfica preta ou azul.",
                "Mantenha os 4 marcadores pretos (■) dos cantos limpos para leitura óptica."
            ]

        for txt in orient_list:
            p = c_orient.add_paragraph()
            p.add_run(f"• {sanitizar_xml_str(txt)}").font.size = Pt(8.5)
            p.paragraph_format.space_after = Pt(1)

        doc.add_paragraph()

        # 3. RENDERIZAÇÃO DO CONTEÚDO
        if is_discursiva:
            renderizar_conteudo_discursivo_com_caixas(doc, corpo_bruto)
        elif is_pei_doc:
            new_section = doc.add_section(WD_SECTION.CONTINUOUS)
            sectPr = new_section._sectPr
            cols = sectPr.xpath('./w:cols')[0]
            cols.set(qn('w:num'), '2')
            cols.set(qn('w:space'), '420')
            renderizar_conteudo_com_tabelas(doc, corpo_bruto)
        else:
            adicionar_cartao_resposta_fiducial_word(doc, num_total_q, is_pei=False)
            doc.add_paragraph()
            new_section = doc.add_section(WD_SECTION.CONTINUOUS)
            sectPr = new_section._sectPr
            cols = sectPr.xpath('./w:cols')[0]
            cols.set(qn('w:num'), '2')
            cols.set(qn('w:space'), '450')
            renderizar_conteudo_com_tabelas(doc, corpo_bruto)

        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        file_stream = io.BytesIO()
        err_doc = Document()
        err_doc.add_paragraph(f"ERRO NO EXPORTER DE PROVA: {sanitizar_xml_str(str(e))}")
        err_doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

# ==============================================================================
# 6. PLANO PEDAGÓGICO SEMANAL
# ==============================================================================
def gerar_docx_plano_pedagogico_ELITE(titulo_arquivo, dados, info):
    file_stream = io.BytesIO()
    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin, section.bottom_margin = Inches(0.4), Inches(0.4)
        section.left_margin, section.right_margin = Inches(0.5), Inches(0.5)

        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        widths = [Inches(1.1), Inches(3.6), Inches(2.0)]
        for i, w in enumerate(widths): table.columns[i].width = w

        logo_path = "logo_escola.png" if os.path.exists("logo_escola.png") else "logo.png"
        if os.path.exists(logo_path):
            cell_logo = table.cell(0, 0).merge(table.cell(2, 0))
            cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_logo = cell_logo.paragraphs[0]
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(logo_path, width=Inches(0.85))
        
        table.cell(0, 1).paragraphs[0].add_run("ESCOLA MUNICIPAL FLAVIO JOSE SIMOES COSTA").font.bold = True
        
        c_tit = table.cell(0, 2)
        set_cell_background(c_tit, "2962FF")
        p_t = c_tit.paragraphs[0]
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_t = p_t.add_run("PLANO DE ENSINO SEMANAL")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(255, 255, 255)

        table.cell(1, 1).paragraphs[0].add_run("Professor: Ronaldo Gomes").font.size = Pt(9.5)
        table.cell(1, 2).paragraphs[0].add_run(f"Série: {sanitizar_xml_str(str(info.get('ano', '')))}").font.size = Pt(9.5)
        table.cell(2, 1).paragraphs[0].add_run(f"Semana: {sanitizar_xml_str(str(info.get('semana', '')))}").font.size = Pt(9.5)
        table.cell(2, 2).paragraphs[0].add_run(f"Trimestre: {sanitizar_xml_str(str(info.get('trimestre', 'I')))}").font.bold = True

        doc.add_paragraph()

        campos = [
            ("OBJETO DE CONHECIMENTO (EIXO):", "geral"), 
            ("CONTEÚDOS ESPECÍFICOS:", "especificos"), 
            ("OBJETIVOS DE APRENDIZAGEM:", "objetivos"), 
            ("RECURSOS DIDÁTICOS:", "recursos"),
            ("PROCEDIMENTOS METODOLÓGICOS:", "metodologia"), 
            ("AVALIAÇÃO E ACOMPANHAMENTO:", "avaliacao"), 
            ("ESTRATÉGIAS DE ACESSIBILIDADE (DUA/PEI):", "pei")
        ]

        for label, chave in campos:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            
            run_label = p.add_run(f"{label}\n")
            run_label.bold = True
            run_label.font.size = Pt(10.5)
            run_label.font.color.rgb = RGBColor(0, 51, 102)
            
            texto_limpo = sanitizar_xml_str(str(dados.get(chave, ""))).replace("**", "").replace("#", "").strip()
            adicionar_texto_formatado(p, texto_limpo)

        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        file_stream = io.BytesIO()
        err_doc = Document(); err_doc.add_paragraph(f"ERRO NO PLANO: {sanitizar_xml_str(str(e))}"); err_doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

# ==============================================================================
# 7. EXPORTADOR PEI NÍVEL 3 (SENSORIAL / MOTOR / BENTO CARDS NO PAPEL)
# ==============================================================================
def gerar_docx_pei_qualitativa(titulo_doc, conteudo, info):
    file_stream = io.BytesIO()
    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin, section.bottom_margin = Inches(0.4), Inches(0.4)
        section.left_margin, section.right_margin = Inches(0.4), Inches(0.4)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10.5)

        configurar_cabecalho_mestre(doc, info, "AVALIAÇÃO ADAPTADA (NÍVEL 3)", mostrar_nota=False)
        doc.add_paragraph()

        panel_info = doc.add_table(rows=1, cols=1)
        panel_info.style = 'Table Grid'
        cell_info = panel_info.cell(0, 0)
        set_cell_background(cell_info, "F1F5F9")
        p_info = cell_info.paragraphs[0]
        p_info.paragraph_format.space_after = Pt(2)
        r_info_title = p_info.add_run("📋 ORIENTAÇÕES DE MEDIAÇÃO PEDAGÓGICA (PEI NÍVEL 3 - ATIVIDADES IMPRESSAS NO PAPEL):\n")
        r_info_title.bold = True
        r_info_title.font.size = Pt(9.5)
        r_info_title.font.color.rgb = RGBColor(0, 51, 102)
        
        orientacoes = [
            "Atividades adaptadas para execução direta na folha (Pintar, Ligar Colunas, Cobrir Pontilhado, Circular).",
            "Não exige suporte material concreto na sala. Utilize lápis de cor, giz de cera e canetinha.",
            "Assinale o nível de autonomia do estudante nas caixas de registro de cada BOX e na Rubrica Final."
        ]
        for o_txt in orientacoes:
            p_o = cell_info.add_paragraph()
            p_o.paragraph_format.space_after = Pt(1)
            p_o.add_run(f"• {o_txt}").font.size = Pt(8.5)

        doc.add_paragraph()

        linhas = sanitizar_xml_str(str(conteudo)).split('\n')
        
        i = 0
        while i < len(linhas):
            linha = linhas[i].strip()
            if not linha:
                i += 1
                continue

            if any(x in linha.upper() for x in ["ATIVIDADE", "TEMA:", "JORNADA", "AVALIAÇÃO ADAPTADA"]) and "BOX" not in linha.upper():
                if "RUBRICA" in linha.upper():
                    break
                p_act = doc.add_paragraph()
                p_act.paragraph_format.space_before = Pt(10)
                p_act.paragraph_format.space_after = Pt(4)
                r_act = p_act.add_run(linha.replace('**', ''))
                r_act.bold = True
                r_act.font.size = Pt(11)
                r_act.font.color.rgb = RGBColor(0, 51, 102)
                i += 1
                continue

            if "[" in linha and "PROMPT IMAGEM" in linha.upper():
                desc_p = re.sub(r'\[\s*PROMPT IMAGEM:\s*|\s*\]', '', linha, flags=re.IGNORECASE)
                adicionar_box_imagem_word(doc, desc_p)
                i += 1
                continue

            if "BOX" in linha.upper():
                card_table = doc.add_table(rows=2, cols=1)
                card_table.style = 'Table Grid'
                card_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                card_table.columns[0].width = Inches(7.5)

                cell_head = card_table.cell(0, 0)
                set_cell_background(cell_head, "2962FF")
                set_row_height(card_table.rows[0], 20)
                p_head = cell_head.paragraphs[0]
                p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                m_box_title = re.search(r"(\[?BOX\s*\d+\]?.*?)(?:[:\-]\s*|\s+)(.*)", linha, re.IGNORECASE)
                if m_box_title:
                    rotulo_box = m_box_title.group(1).upper().replace("[", "").replace("]", "").strip()
                    desc_box = m_box_title.group(2).strip()
                else:
                    rotulo_box = "BOX DE ATIVIDADE IMPRESSA"
                    desc_box = linha.strip()

                r_head = p_head.add_run(f"📦 {rotulo_box}")
                r_head.bold = True
                r_head.font.size = Pt(10)
                r_head.font.color.rgb = RGBColor(255, 255, 255)

                cell_body = card_table.cell(1, 0)
                set_cell_background(cell_body, "F8FAFC")
                p_body = cell_body.paragraphs[0]
                p_body.paragraph_format.space_after = Pt(4)
                
                adicionar_texto_formatado(p_body, desc_box)

                p_check = cell_body.add_paragraph()
                p_check.paragraph_format.space_before = Pt(4)
                p_check.paragraph_format.space_after = Pt(2)
                
                r_check_label = p_check.add_run("Registro do Mediador no Papel: ")
                r_check_label.bold = True
                r_check_label.font.size = Pt(8.5)
                r_check_label.font.color.rgb = RGBColor(100, 116, 139)

                r_opts = p_check.add_run("[   ] Autônomo   |   [   ] Com Apoio   |   [   ] Não Realizou")
                r_opts.font.size = Pt(8.5)
                r_opts.font.bold = True
                r_opts.font.color.rgb = RGBColor(30, 41, 59)

                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_after = Pt(2)
                i += 1
                continue

            p_norm = doc.add_paragraph()
            p_norm.paragraph_format.space_after = Pt(4)
            adicionar_texto_formatado(p_norm, linha)
            i += 1

        doc.add_paragraph()
        p_rub_title = doc.add_paragraph()
        p_rub_title.paragraph_format.space_before = Pt(10)
        p_rub_title.paragraph_format.space_after = Pt(4)
        r_rub = p_rub_title.add_run("📋 RUBRICA OFICIAL DE OBSERVAÇÃO PEDAGÓGICA (PARECER QUALITATIVO)")
        r_rub.bold = True
        r_rub.font.size = Pt(11)
        r_rub.font.color.rgb = RGBColor(0, 51, 102)

        rubrica_table = doc.add_table(rows=5, cols=5)
        rubrica_table.style = 'Table Grid'
        
        col_widths = [Inches(2.45), Inches(1.1), Inches(1.1), Inches(1.15), Inches(1.6)]
        for row in rubrica_table.rows:
            set_row_height(row, 22)
            for idx_c, w in enumerate(col_widths):
                row.cells[idx_c].width = w

        headers_rub = ["DIMENSÃO COGNITIVA / MOTORA", "AUTÔNOMO ✅", "COM APOIO 🤝", "NÃO REALIZOU ❌", "OBSERVAÇÕES"]
        for idx_c, h_text in enumerate(headers_rub):
            c_h = rubrica_table.cell(0, idx_c)
            set_cell_background(c_h, "003366")
            p_h = c_h.paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_h = p_h.add_run(h_text)
            r_h.bold = True
            r_h.font.size = Pt(8.5)
            r_h.font.color.rgb = RGBColor(255, 255, 255)

        dimensoes_pei = [
            ("Autonomia Executiva", "Iniciativa e condução motora nas tarefas de papel"),
            ("Compreensão de Comandos", "Atendimento às instruções de pintar, ligar e cobrir"),
            ("Percepção Visual e Espacial", "Identificação de formas, números e sequências no papel"),
            ("Raciocínio Lógico-Proporcional", "Associação visual de quantidades e símbolos")
        ]

        for idx_d, (dim_nome, dim_desc) in enumerate(dimensoes_pei, start=1):
            row_cells = rubrica_table.rows[idx_d].cells
            
            p_dim = row_cells[0].paragraphs[0]
            r_dn = p_dim.add_run(f"{dim_nome}\n")
            r_dn.bold = True
            r_dn.font.size = Pt(9)
            r_dd = p_dim.add_run(dim_desc)
            r_dd.font.size = Pt(7.5)
            r_dd.font.italic = True
            r_dd.font.color.rgb = RGBColor(100, 116, 139)

            for c_idx in range(1, 4):
                p_chk = row_cells[c_idx].paragraphs[0]
                p_chk.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_chk.add_run("○").font.size = Pt(14)

            row_cells[4].paragraphs[0].add_run("____________________").font.size = Pt(8)

        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        file_stream = io.BytesIO()
        err_doc = Document(); err_doc.add_paragraph(f"ERRO NO EXPORTER PEI N3: {sanitizar_xml_str(str(e))}"); err_doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

# ==============================================================================
# 8. ETIQUETAS DE NOTAS (CALIBRADA PARA A4 - TOTAL 7.50 IN)
# ==============================================================================
def gerar_docx_etiquetas_notas(nome_arquivo, dados_alunos, info):
    """
    SOSA V2026.PRO_INFINITY - ETIQUETAS OFICIAIS DE NOTAS, REFACÇÃO (+0.5) E RECUPERAÇÃO
    Design executivo A4 em 2 colunas perfeitamente calibrado para folha de papel A4.
    """
    file_stream = io.BytesIO()
    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = section.bottom_margin = Inches(0.35)
        section.left_margin = section.right_margin = Inches(0.35)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(9)

        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Inches(3.75)
        table.columns[1].width = Inches(3.75)

        turma_label = sanitizar_xml_str(str(info.get('turma', '6º Ano')))
        trim_label = sanitizar_xml_str(str(info.get('trimestre', 'II Trimestre')))

        for i in range(0, len(dados_alunos), 2):
            row = table.add_row()
            set_row_height(row, 130)

            for j in range(2):
                if i + j < len(dados_alunos):
                    aluno = dados_alunos[i + j]
                    c = row.cells[j]
                    c.width = Inches(3.75)
                    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    set_cell_background(c, "FAFAFA")
                    
                    p = c.paragraphs[0]
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.1
                    
                    r_esc = p.add_run("ESCOLA MUNICIPAL FLÁVIO JOSÉ SIMÕES COSTA\n")
                    r_esc.bold = True
                    r_esc.font.size = Pt(8.5)
                    r_esc.font.color.rgb = RGBColor(0, 51, 102)
                    
                    nome_al_limpo = sanitizar_xml_str(str(aluno.get('nome', 'Estudante'))).replace("👤 ", "").replace("♿ ", "").replace("🟠 ", "").replace("🧱 ", "").replace("🧮 ", "").replace("🚀 ", "")
                    p.add_run(f"Estudante: {nome_al_limpo}\n").bold = True
                    p.add_run(f"Turma: {turma_label} | Período: {trim_label}\n").font.size = Pt(8.0)
                    
                    c1_v = sanitizar_xml_str(str(aluno.get('c1', aluno.get('vistos', '0.0'))))
                    c2_v = sanitizar_xml_str(str(aluno.get('c2', aluno.get('teste', '0.0'))))
                    c3_v = sanitizar_xml_str(str(aluno.get('c3', aluno.get('prova', '0.0'))))
                    bonus_v = sanitizar_xml_str(str(aluno.get('bonus', '0.0')))
                    media_v = sanitizar_xml_str(str(aluno.get('media', '0.0')))
                    media_num = helper_sosa_float(media_v)

                    p_comp = p.add_run(f"Caderno (C1): {c1_v} | Testes (C2): {c2_v} | Prova (C3): {c3_v} | Bônus: {bonus_v}\n")
                    p_comp.font.size = Pt(8.0)
                    p_comp.font.color.rgb = RGBColor(71, 85, 105)
                    
                    r_med = p.add_run(f"MÉDIA DO TRIMESTRE: {media_num:.1f} pontos\n")
                    r_med.bold = True
                    r_med.font.size = Pt(9.5)
                    
                    if media_num >= 6.0:
                        r_st = p.add_run("SITUAÇÃO: APROVADO NO TRIMESTRE\n")
                        r_st.bold = True
                        r_st.font.size = Pt(8.5)
                        r_st.font.color.rgb = RGBColor(0, 128, 0)
                        p.add_run("• Parabéns! Você atingiu a média e está dispensado da recuperação.\n").font.size = Pt(8.0)
                        r_ref = p.add_run("• Oportunidade de Refacção: Entregue a prova corrigida no caderno para somar +0.5 pts e elevar sua média!\n")
                        r_ref.font.size = Pt(7.5)
                        r_ref.font.italic = True
                        
                    elif media_num == 5.5:
                        r_st = p.add_run("SITUAÇÃO: OPORTUNIDADE DE REFACÇÃO SOLIDÁRIA (+0.5)\n")
                        r_st.bold = True
                        r_st.font.size = Pt(8.5)
                        r_st.font.color.rgb = RGBColor(204, 102, 0)
                        p.add_run("• Refacção no Caderno: Refaça as questões que errou no caderno para somar +0.5 pts, atingir a MÉDIA 6.0 e ser dispensado da prova!\n").bold = True
                        r_rec_info = p.add_run("• Prova de Recuperação (Opcional): Caso prefira fazer a prova (0 a 10), precisará de 6.5. Regra: (Média + Prova) ÷ 2 ≥ 6.0\n")
                        r_rec_info.font.size = Pt(7.5)
                        r_rec_info.font.italic = True
                        
                    else:
                        r_st = p.add_run("SITUAÇÃO: CONVOCADO PARA RECUPERAÇÃO PARALELA\n")
                        r_st.bold = True
                        r_st.font.size = Pt(8.5)
                        r_st.font.color.rgb = RGBColor(192, 0, 0)
                        
                        nota_necessaria_rec = max(0.0, 12.0 - media_num)
                        if nota_necessaria_rec <= 10.0:
                            p.add_run(f"• Meta na Prova de Recuperação (0 a 10): Tirar no mínimo {nota_necessaria_rec:.1f} pontos\n").bold = True
                        else:
                            p.add_run("• Meta na Prova de Recuperação: Tirar 10.0 pontos (Necessita de Conselho Final)\n").bold = True
                            
                        p.add_run("• Refacção Solidária: Entregue a prova corrigida no caderno para somar +0.5 pts na sua base!\n").font.size = Pt(7.5)
                        r_regra = p.add_run("• Regra Oficial: (Média do Trimestre + Prova de Recuperação) ÷ 2 ≥ 6.0\n")
                        r_regra.font.size = Pt(7.5)
                        r_regra.font.italic = True

                    r_prof = p.add_run("Prof. Ronaldo Gomes • Componente Curricular de Matemática")
                    r_prof.font.size = Pt(7.5)
                    r_prof.font.color.rgb = RGBColor(100, 116, 139)

        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        file_stream = io.BytesIO()
        err_doc = Document()
        err_doc.add_paragraph(f"ERRO NAS ETIQUETAS: {sanitizar_xml_str(str(e))}")
        err_doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

# ==============================================================================
# 9. PLANEJAMENTO TRIMESTRAL (PAISAGEM)
# ==============================================================================
def gerar_docx_planejamento_trimestral(nome_arquivo, info, dados_tabela):
    file_stream = io.BytesIO()
    try:
        doc = Document()
        section = doc.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = section.right_margin = Inches(0.4)
        section.top_margin = section.bottom_margin = Inches(0.4)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(9)

        p_cab = doc.add_paragraph()
        p_cab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_esc = p_cab.add_run("ESCOLA MUNICIPAL FLÁVIO JOSÉ SIMÕES COSTA\n")
        run_esc.bold = True
        run_esc.font.size = Pt(11)
        
        run_tit = p_cab.add_run(f"PLANEJAMENTO TRIMESTRAL DE MATEMÁTICA - {sanitizar_xml_str(str(info['trimestre'])).upper()} / 2026\n")
        run_tit.bold = True
        run_tit.font.size = Pt(10)
        run_sub = p_cab.add_run(f"SÉRIE: {sanitizar_xml_str(str(info['ano']))}   |   PROFESSOR: RONALDO GOMES")
        run_sub.font.size = Pt(9)
        
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        widths = [Inches(1.8), Inches(3.8), Inches(2.2), Inches(3.0)]
        for i, w in enumerate(widths): table.columns[i].width = w

        headers = ['EIXO TEMÁTICO', 'CONTEÚDOS PROGRAMÁTICOS', 'HABILIDADES (BNCC)', 'METODOLOGIA APLICADA']
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            set_cell_background(cell, "003366")
            cell.text = h
            p = cell.paragraphs[0]
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        for row_data in dados_tabela:
            row_cells = table.add_row().cells
            row_cells[0].text = sanitizar_xml_str(str(row_data['eixo']))
            row_cells[1].text = sanitizar_xml_str(str(row_data['conteudos']))
            row_cells[2].text = sanitizar_xml_str(str(row_data['habilidades']))
            row_cells[3].text = sanitizar_xml_str(str(row_data['metodologia']))
            
            for cell in row_cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        file_stream = io.BytesIO()
        err_doc = Document(); err_doc.add_paragraph(f"ERRO NO PLANO TRIMESTRAL: {sanitizar_xml_str(str(e))}"); err_doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

# ==============================================================================
# 10. PEI OFICIAL DA PREFEITURA DE ITABUNA (FOCO: COMPONENTE DE MATEMÁTICA)
# ==============================================================================
def gerar_docx_pei_oficial(nome_arquivo, dados_aluno, habilidades, curriculo_df, parecer_resultados=""):
    """
    SOSA V2026 - MODELO OFICIAL DA SECRETARIA MUNICIPAL DA EDUCAÇÃO DE ITABUNA:
    Coordenação Técnica Pedagógica da Educação Especial.
    Gera o documento institucional com foco no Componente de MATEMÁTICA:
    - Página 1: Cabeçalho Oficial, Dados da Estudante e Programas.
    - Páginas 2-3: Estudo de Caso (Sociais, Comunicativas, Emocionais, Funcionais) e Medida de Acesso.
    - Páginas 4+: Tabela de Planejamento de Matemática (I, II e III Trimestres).
    - Página Final: Resultados Obtidos em Matemática / Ciências e Assinaturas Oficiais.
    """
    file_stream = io.BytesIO()
    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(9.5)

        p_cab = doc.add_paragraph()
        p_cab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cab.paragraph_format.space_after = Pt(2)
        r_cab1 = p_cab.add_run("SECRETARIA MUNICIPAL DA EDUCAÇÃO\nDEPARTAMENTO DE EDUCAÇÃO BÁSICA\nCOORDENAÇÃO TÉCNICA PEDAGÓGICA DA EDUCAÇÃO ESPECIAL\n")
        r_cab1.bold = True
        r_cab1.font.size = Pt(10)
        
        p_tit = doc.add_paragraph()
        p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_tit.paragraph_format.space_before = Pt(6)
        p_tit.paragraph_format.space_after = Pt(4)
        r_tit = p_tit.add_run("PLANO EDUCACIONAL INDIVIDUALIZADO - PEI\n")
        r_tit.bold = True
        r_tit.font.size = Pt(11.5)
        
        p_obj = doc.add_paragraph()
        p_obj.paragraph_format.space_after = Pt(8)
        r_obj_lbl = p_obj.add_run("OBJETIVO: ")
        r_obj_lbl.bold = True
        p_obj.add_run("Planejar, adaptar e implementar estratégias pedagógicas garantindo acessibilidade curricular do aluno na Unidade Escolar.").font.size = Pt(9)

        p_d_title = doc.add_paragraph()
        p_d_title.paragraph_format.space_after = Pt(2)
        p_d_title.add_run("DADOS DO ALUNO").bold = True

        tb_dados = doc.add_table(rows=6, cols=2)
        tb_dados.style = 'Table Grid'
        tb_dados.columns[0].width = Inches(4.5)
        tb_dados.columns[1].width = Inches(2.9)
        for r in tb_dados.rows: set_row_height(r, 18)

        nome_al = sanitizar_xml_str(str(dados_aluno.get('nome', 'ESTUDANTE'))).upper()
        turma_al = sanitizar_xml_str(str(dados_aluno.get('turma', '6MA'))).upper()
        cid_al = sanitizar_xml_str(str(dados_aluno.get('cid', 'F84'))).upper()
        idade_al = sanitizar_xml_str(str(dados_aluno.get('idade', '11')))
        ano_letivo = "2026"

        tb_dados.cell(0, 0).paragraphs[0].add_run("UNIDADE ESCOLAR: Escola Municipal Flávio José Simões Costa").font.size = Pt(8.5)
        tb_dados.cell(0, 1).paragraphs[0].add_run(f"ANO LETIVO: {ano_letivo}").font.size = Pt(8.5)

        tb_dados.cell(1, 0).paragraphs[0].add_run(f"NOME: {nome_al}").font.size = Pt(8.5)
        tb_dados.cell(1, 1).paragraphs[0].add_run(f"TURMA: {turma_al} | IDADE: {idade_al} anos").font.size = Pt(8.5)

        tb_dados.cell(2, 0).paragraphs[0].add_run("Nome do Responsável: Conforme Cadastro Escolar").font.size = Pt(8.5)
        tb_dados.cell(2, 1).paragraphs[0].add_run("Fone: (73) 9XXXX-XXXX").font.size = Pt(8.5)

        tb_dados.cell(3, 0).merge(tb_dados.cell(3, 1))
        tb_dados.cell(3, 0).paragraphs[0].add_run("PARTICIPA DE PROJETOS/PROGRAMAS: Geia ( )  SRM ( )  CEPEI ( )  Psicólogo ( )  Psiquiatra ( )  CAPS ( )  CREADE ( )  Outros ( )").font.size = Pt(8.0)

        tb_dados.cell(4, 0).paragraphs[0].add_run(f"DEFICIÊNCIA(S) / CID: {cid_al}").font.size = Pt(8.5)
        tb_dados.cell(4, 1).paragraphs[0].add_run("SUSPEIÇÃO: ( )").font.size = Pt(8.5)

        tb_dados.cell(5, 0).paragraphs[0].add_run("NÍVEL DE SUPORTE: Nível de Acessibilidade Curricular").font.size = Pt(8.5)
        tb_dados.cell(5, 1).paragraphs[0].add_run("MONITOR: ( )").font.size = Pt(8.5)

        doc.add_paragraph()

        p_est_tit = doc.add_paragraph()
        p_est_tit.paragraph_format.space_before = Pt(6)
        p_est_tit.paragraph_format.space_after = Pt(2)
        r_est = p_est_tit.add_run("1 - Plano de acessibilidade curricular.")
        r_est.bold = True
        r_est.font.size = Pt(10.5)

        p_obs_intro = doc.add_paragraph()
        p_obs_intro.paragraph_format.space_after = Pt(4)
        p_obs_intro.add_run("Com base no estudo de caso: sondagem, observação em sala de aula de Matemática e devolutiva trimestral, observa-se que a estudante apresenta necessidades em:").font.size = Pt(8.5)

        dimensoes_oficiais = [
            ("Habilidades Sociais", "comportamentos repetitivos e restritos ( X )  estereotipias ( X )  níveis de brincadeiras ( )  rotina ( X )  isolamento ( )  atenção compartilhada ( X )  outras ( )", "Habilidades Sociais"),
            ("Habilidades Comunicativas", "comunicação verbal ( )  comunicação não verbal ( X )  clareza de comunicação ( )  contato visual ( X )  toque físico ( )  compreensão na comunicação ( X )  comunicação alternativa ( X )  linguagem expressiva e receptiva ( X )", "Habilidades Comunicativas"),
            ("Habilidades Emocionais", "controle inibitório ( X )  resposta emocional ( )  flexibilidade ( X )  empatia ( )", "Habilidades Emocionais"),
            ("Habilidades Funcionais", "Auto cuidado ( )  higiene pessoal ( )  alimentação ( )  vestimentas ( )  organização de material escolar ( X )", "Habilidades Funcionais")
        ]

        for dim_titulo, dim_checks, dim_key in dimensoes_oficiais:
            p_dim = doc.add_paragraph()
            p_dim.paragraph_format.space_before = Pt(4)
            p_dim.paragraph_format.space_after = Pt(1)
            p_dim.add_run(f"• {dim_titulo}:\n").bold = True
            r_chk = p_dim.add_run(f"{dim_checks}\n")
            r_chk.font.size = Pt(8.0)
            r_chk.font.color.rgb = RGBColor(71, 85, 105)

            tb_desc = doc.add_table(rows=1, cols=1)
            tb_desc.style = 'Table Grid'
            tb_desc.columns[0].width = Inches(7.4)
            set_row_height(tb_desc.rows[0], 28)
            c_desc = tb_desc.cell(0, 0)
            set_cell_background(c_desc, "FAFAFA")
            
            p_c = c_desc.paragraphs[0]
            p_c.paragraph_format.space_before = Pt(2)
            p_c.paragraph_format.space_after = Pt(2)
            
            texto_hab = sanitizar_xml_str(str(habilidades.get(dim_key, 'Observação mediada nas atividades de sala de aula.'))).strip()
            if not texto_hab: texto_hab = "Necessidade de mediação e suporte visual durante as tarefas."
            p_c.add_run(f"Descrição: {texto_hab}").font.size = Pt(8.5)

        doc.add_paragraph()
        p_medida = doc.add_paragraph()
        p_medida.paragraph_format.space_before = Pt(6)
        p_medida.paragraph_format.space_after = Pt(2)
        p_medida.add_run("2 - Plano Trimestral (Medidas de acesso ao currículo conforme o destinatário):\n").bold = True
        p_medida.add_run("(   ) Currículo Funcional (Ativ. de Vida Diária)     ( X ) Currículo Adaptado (Objetivos Específicos)     (   ) Currículo Suplementar").font.size = Pt(8.5)

        doc.add_paragraph()

        p_plan_title = doc.add_paragraph()
        p_plan_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_plan_title.paragraph_format.space_before = Pt(8)
        p_plan_title.paragraph_format.space_after = Pt(4)
        r_ptit = p_plan_title.add_run("PLANEJAMENTO POR COMPONENTE CURRICULAR — MATEMÁTICA")
        r_ptit.bold = True
        r_ptit.font.size = Pt(11)
        r_ptit.font.color.rgb = RGBColor(0, 51, 102)

        tb_curr = doc.add_table(rows=1, cols=4)
        tb_curr.style = 'Table Grid'
        
        col_w = [Inches(1.3), Inches(2.3), Inches(2.1), Inches(1.7)]
        for i_w, w in enumerate(col_w): tb_curr.columns[i_w].width = w

        hdr_cells = tb_curr.rows[0].cells
        hdr_cells[0].text = 'DISCIPLINA'
        hdr_cells[1].text = 'OBJETIVOS APRENDIZAGEM'
        hdr_cells[2].text = 'ESTRATÉGIAS METODOLÓGICAS'
        hdr_cells[3].text = 'RECURSOS MATERIAIS'

        set_row_height(tb_curr.rows[0], 20)
        for cell in hdr_cells:
            set_cell_background(cell, "003366")
            p_h = cell.paragraphs[0]
            p_h.runs[0].bold = True
            p_h.runs[0].font.size = Pt(8.5)
            p_h.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def limpar_coluna_estrita(texto, tipo='OBJETIVO'):
            if not texto or not isinstance(texto, str): return ""
            t = str(texto).strip()
            if tipo == 'OBJETIVO':
                t = re.split(r'\[ESTRATEGIA\]|\[RECURSO\]|\[/ITEM\]', t, flags=re.IGNORECASE)[0]
            elif tipo == 'ESTRATEGIA':
                t = re.split(r'\[RECURSO\]|\[/ITEM\]', t, flags=re.IGNORECASE)[0]
            elif tipo == 'RECURSO':
                t = re.split(r'\[/ITEM\]|\[ITEM\]', t, flags=re.IGNORECASE)[0]
            t = re.sub(r'\[/?(?:ITEM|OBJETIVO|ESTRATEGIA|RECURSO)\]', '', t, flags=re.IGNORECASE).strip()
            return sanitizar_xml_str(t)

        if not curriculo_df.empty:
            for _, row_c in curriculo_df.iterrows():
                row_cells = tb_curr.add_row().cells
                set_row_height(tb_curr.rows[-1], 24)
                
                row_cells[0].text = "MATEMÁTICA\n(Prof. Ronaldo Gomes)"
                row_cells[1].text = limpar_coluna_estrita(row_c.get('Objetivos de Aprendizagem', ''), 'OBJETIVO')
                row_cells[2].text = limpar_coluna_estrita(row_c.get('Estratégias Metodológicas', ''), 'ESTRATEGIA')
                row_cells[3].text = limpar_coluna_estrita(row_c.get('Recursos Materiais', ''), 'RECURSO')
                
                for idx_c, cell in enumerate(row_cells):
                    cell.width = col_w[idx_c]
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        if idx_c == 0:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.runs[0].font.bold = True
                            p.runs[0].font.size = Pt(8.0)
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p.runs[0].font.size = Pt(8.0)
        else:
            row_cells = tb_curr.add_row().cells
            row_cells[0].text = "MATEMÁTICA"
            row_cells[1].text = "Desenvolver o raciocínio lógico e a compreensão de quantidades através de representação visual e materiais concretos."
            row_cells[2].text = "Instruções curtas e estruturadas em pequenas etapas; mediação individualizada nas tarefas; uso de cores para diferenciação."
            row_cells[3].text = "Material Dourado, malha quadriculada, ábaco, calculadora, atividades adaptadas impressas em papel ofício."

        doc.add_paragraph()

        p_res_title = doc.add_paragraph()
        p_res_title.paragraph_format.space_before = Pt(8)
        p_res_title.paragraph_format.space_after = Pt(2)
        r_rtit = p_res_title.add_run("3 - Resultados obtidos diante dos objetivos do Plano Educacional Individualizado - PEI\n(descrição por área de conhecimento):")
        r_rtit.bold = True
        r_rtit.font.size = Pt(9.5)

        tb_res = doc.add_table(rows=3, cols=1)
        tb_res.style = 'Table Grid'
        tb_res.columns[0].width = Inches(7.4)

        c_r1 = tb_res.cell(0, 0)
        c_r1.paragraphs[0].add_run("PORTUGUÊS / ED. FÍSICA / ARTES\n").bold = True
        c_r1.paragraphs[0].runs[0].font.size = Pt(8.5)
        c_r1.paragraphs[0].add_run("(A cargo dos docentes da área)").font.size = Pt(8.0)

        c_r2 = tb_res.cell(1, 0)
        set_cell_background(c_r2, "F8FAFC")
        p_mat_res = c_r2.paragraphs[0]
        r_m_lbl = p_mat_res.add_run("MATEMÁTICA / CIÊNCIAS\n")
        r_m_lbl.bold = True
        r_m_lbl.font.size = Pt(8.5)
        r_m_lbl.font.color.rgb = RGBColor(0, 51, 102)

        parecer_mat_txt = sanitizar_xml_str(str(parecer_resultados)).strip()
        if not parecer_mat_txt:
            parecer_mat_txt = "Matemática: A estudante encontra-se em pleno processo de desenvolvimento da aprendizagem, respondendo positivamente ao uso de materiais manipulativos concretos e à mediação individualizada, demonstrando evolução na compreensão dos conceitos e na autonomia para a resolução de atividades."
        p_mat_res.add_run(parecer_mat_txt).font.size = Pt(8.5)

        c_r3 = tb_res.cell(2, 0)
        c_r3.paragraphs[0].add_run("GEOGRAFIA / HISTÓRIA\n").bold = True
        c_r3.paragraphs[0].runs[0].font.size = Pt(8.5)
        c_r3.paragraphs[0].add_run("(A cargo dos docentes da área)").font.size = Pt(8.0)

        for r in tb_res.rows: set_row_height(r, 32)

        doc.add_paragraph()

        hoje_dia = datetime.now().strftime("%d")
        meses_pt = {"01":"Janeiro","02":"Fevereiro","03":"Março","04":"Abril","05":"Maio","06":"Junho","07":"Julho","08":"Agosto","09":"Setembro","10":"Outubro","11":"Novembro","12":"Dezembro"}
        mes_atual = meses_pt.get(datetime.now().strftime("%m"), "Agosto")
        
        p_data = doc.add_paragraph()
        p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_data.paragraph_format.space_before = Pt(6)
        p_data.add_run(f"Itabuna, {hoje_dia} de {mes_atual} de 2026.").font.size = Pt(9)

        doc.add_paragraph()

        p_ass = doc.add_paragraph()
        p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ass.paragraph_format.space_after = Pt(2)
        p_ass.paragraph_format.line_spacing = 1.3
        
        p_ass.add_run("________________________________________________\nProfessor Linguagem (Língua Portuguesa, Arte e Educação Física)\n\n").font.size = Pt(8.0)
        
        r_ron = p_ass.add_run("_________________________________________________\nProfessor Ciências da Natureza e Matemática: Prof. Ronaldo Gomes\n\n")
        r_ron.bold = True
        r_ron.font.size = Pt(8.5)
        r_ron.font.color.rgb = RGBColor(0, 51, 102)
        
        p_ass.add_run("_________________________________________________\nProfessor Ciências Humanas (História e Geografia)\n\n").font.size = Pt(8.0)
        p_ass.add_run("__________________________________________________\nCoordenador Pedagógico").font.size = Pt(8.0)

        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        file_stream = io.BytesIO()
        err_doc = Document()
        err_doc.add_paragraph(f"ERRO NO PEI OFICIAL: {sanitizar_xml_str(str(e))}")
        err_doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

# ==============================================================================
# 11. CERTIDÃO OFICIAL DE PRODUÇÃO E RENDIMENTO
# ==============================================================================
def gerar_docx_certidao_producao(nome_arquivo, dados_aluno, notas_trimestres, info_escola):
    """
    SOSA V2026: CERTIDÃO OFICIAL DE PRODUÇÃO E RENDIMENTO DO ESTUDANTE INATIVO/TRANSFERIDO.
    Gera um documento Word A4 assinado e formatado para a Direção/Coordenação.
    """
    file_stream = io.BytesIO()
    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.5)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        configurar_cabecalho_mestre(doc, info_escola, "CERTIDÃO DE RENDIMENTO E PRODUÇÃO", mostrar_nota=False)
        doc.add_paragraph()

        p_tit = doc.add_paragraph()
        p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_tit.add_run("CERTIDÃO OFICIAL DE RENDIMENTO E PRODUÇÃO PEDAGÓGICA")
        run_t.bold = True
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()

        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_after = Pt(4)
        p_d.add_run("ESTUDANTE: ").bold = True
        p_d.add_run(f"{sanitizar_xml_str(str(dados_aluno.get('nome', '')))}\t\t")
        p_d.add_run("MATRÍCULA/ID: ").bold = True
        p_d.add_run(f"{sanitizar_xml_str(str(dados_aluno.get('id', '')))}\n")
        
        p_d2 = doc.add_paragraph()
        p_d2.paragraph_format.space_after = Pt(6)
        p_d2.add_run("TURMA: ").bold = True
        p_d2.add_run(f"{sanitizar_xml_str(str(dados_aluno.get('turma', '')))}\t\t")
        p_d2.add_run("STATUS REGIMENTAL: ").bold = True
        
        status_txt = sanitizar_xml_str(str(dados_aluno.get('status', 'TRANSFERIDO'))).upper()
        r_st = p_d2.add_run(status_txt)
        r_st.bold = True
        if "ATIVO" in status_txt: r_st.font.color.rgb = RGBColor(0, 128, 0)
        else: r_st.font.color.rgb = RGBColor(204, 0, 0)

        p_d2.add_run(f"\t\tPERFIL/CID: {sanitizar_xml_str(str(dados_aluno.get('perfil', 'TÍPICO')))}")

        doc.add_paragraph()

        p_tb_t = doc.add_paragraph()
        p_tb_t.add_run("1. HISTÓRICO DE NOTAS E AVALIAÇÕES ACUMULADAS:").bold = True
        p_tb_t.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        table_n = doc.add_table(rows=1, cols=6)
        table_n.style = 'Table Grid'
        hdr_cells = table_n.rows[0].cells
        hdr_titles = ["PERÍODO", "C1 (VISTOS)", "C2 (TESTES)", "C3 (PROVA)", "REC.", "MÉDIA"]
        col_w = [Inches(1.5), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.0), Inches(1.2)]
        
        set_row_height(table_n.rows[0], 20)
        for idx_h, h_t in enumerate(hdr_titles):
            c = hdr_cells[idx_h]
            c.width = col_w[idx_h]
            set_cell_background(c, "003366")
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h_t)
            r.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

        for reg_trim in notas_trimestres:
            row_cells = table_n.add_row().cells
            set_row_height(table_n.rows[-1], 18)
            vals = [
                reg_trim.get('periodo', ''),
                f"{helper_sosa_float(reg_trim.get('c1', 0.0)):.1f}",
                f"{helper_sosa_float(reg_trim.get('c2', 0.0)):.1f}",
                f"{helper_sosa_float(reg_trim.get('c3', 0.0)):.1f}",
                reg_trim.get('rec', '-'),
                f"{helper_sosa_float(reg_trim.get('media', 0.0)):.1f}"
            ]
            for idx_v, val in enumerate(vals):
                c = row_cells[idx_v]
                c.width = col_w[idx_v]
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(sanitizar_xml_str(str(val))).font.size = Pt(9)

        doc.add_paragraph()

        p_ass = doc.add_paragraph()
        p_ass.add_run("2. REGISTRO DE FREQUÊNCIA E ENGAJAMENTO EM SALA:\n").bold = True
        p_ass.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        p_ass.add_run(f"• Percentual de Assiduidade: {dados_aluno.get('assiduidade', '100%')}\n")
        p_ass.add_run(f"• Total de Ausências Registradas: {dados_aluno.get('faltas', 0)} falta(s)\n")
        p_ass.add_run(f"• Engajamento em Tarefas de Caderno (Vistos C1): {dados_aluno.get('vistos_perc', '0%')}\n")
        p_ass.add_run(f"• Bônus Atitudinais Conquistados: {dados_aluno.get('bonus', '0.0')} pts\n")

        doc.add_paragraph()

        p_par = doc.add_paragraph()
        p_par.add_run("3. PARECER DESCRITIVO DE REGÊNCIA:\n").bold = True
        p_par.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        p_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        parecer_txt = dados_aluno.get('parecer', 'O estudante cumpriu as atividades regulares durante o período em que esteve vinculado à turma.')
        adicionar_texto_formatado(p_par, parecer_txt)

        doc.add_paragraph()
        doc.add_paragraph()

        p_sig = doc.add_paragraph()
        p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sig.add_run("_________________________________________\n").bold = True
        p_sig.add_run("Prof. Ronaldo Gomes\n").bold = True
        p_sig.add_run("Componente Curricular de Matemática\n").font.size = Pt(8.5)
        p_sig.add_run(f"Emitido em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}").font.size = Pt(8.5)

        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        file_stream = io.BytesIO()
        err_doc = Document()
        err_doc.add_paragraph(f"ERRO NA CERTIDÃO: {sanitizar_xml_str(str(e))}")
        err_doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
